from pathlib import Path
from typing import Any, Dict, List, Tuple
import hashlib
import re

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader

try:
    from docx import Document as DocxDocument  # type: ignore

    _HAS_DOCX = True
except Exception:
    DocxDocument = None  # type: ignore
    _HAS_DOCX = False

from app.core.config import CHUNK_SIZE, CHUNK_OVERLAP

E5_PASSAGE_PREFIX = "passage: "
_DOC_ID_PATTERN = re.compile(r"^(\d+)_(.+)$")
_ARTICLE_NUMBER_PATTERN = re.compile(r"^[0-9\u06F0-\u06F9]+")


def _content_hash(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text.strip())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def hash_page_content(page_content: str) -> str:
    """Compute content_hash for stored or chunked text (strips E5 passage prefix)."""
    text = page_content.strip()
    if text.startswith(E5_PASSAGE_PREFIX):
        text = text[len(E5_PASSAGE_PREFIX) :]
    return _content_hash(text)


def _e5_passage_text(text: str) -> str:
    """Prefix document text for multilingual-e5 passage encoding."""
    stripped = text.strip()
    if stripped.startswith(E5_PASSAGE_PREFIX):
        return stripped
    return f"{E5_PASSAGE_PREFIX}{stripped}"


def parse_source_metadata(source: str) -> Dict[str, Any]:
    """Extract doc_id and law_name from ghavanin-style filenames."""
    stem = Path(source).stem
    match = _DOC_ID_PATTERN.match(stem)
    if match:
        return {
            "doc_id": match.group(1),
            "law_name": match.group(2).strip(),
        }
    return {"law_name": stem}


def _extract_article_number(unit_kind: str, unit_title: str) -> str | None:
    if unit_kind not in {"ماده", "اصل", "تبصره", "بند"}:
        return None
    number_match = _ARTICLE_NUMBER_PATTERN.match(unit_title.strip())
    if number_match:
        return number_match.group(0)
    return None


def _build_chunk_metadata(
    source: str,
    document_type: str,
    legal_domain: str,
    content: str,
    *,
    unit_kind: str | None = None,
    unit_title: str | None = None,
    unit_index: int | None = None,
    domain: str | None = None,
    subdomain: str | None = None,
) -> Dict[str, Any]:
    metadata: Dict[str, Any] = {
        "source": source,
        "document_type": document_type,
        "legal_domain": legal_domain,
        "domain": domain or "نامشخص",
        "content_hash": _content_hash(content),
        **parse_source_metadata(source),
    }
    if subdomain:
        metadata["subdomain"] = subdomain
    if unit_kind is not None:
        metadata["unit_kind"] = unit_kind
    if unit_title is not None:
        metadata["unit_title"] = unit_title
        article_number = _extract_article_number(unit_kind or "", unit_title)
        if article_number:
            metadata["article_number"] = article_number
    if unit_index is not None:
        metadata["unit_index"] = unit_index
    return metadata


def _tag_taxonomy(source: str, text: str) -> Dict[str, Any]:
    """Attach hierarchical taxonomy labels (heuristic; optional LLM)."""
    from app.core.config import USE_LLM_TAXONOMY_TAGGING, OPENAI_API_KEY
    from app.services.taxonomy import heuristic_tag_text, taxonomy_prompt_text

    tag = heuristic_tag_text(source, text)
    if (
        USE_LLM_TAXONOMY_TAGGING
        and OPENAI_API_KEY
        and (tag.get("domain") == "نامشخص" or float(tag.get("confidence") or 0) < 0.55)
    ):
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.messages import SystemMessage, HumanMessage
            import json as _json

            snippet = (text or "")[:1800]
            llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
            resp = llm.invoke(
                [
                    SystemMessage(
                        content=(
                            "متن قانونی زیر را بر اساس این taxonomy طبقه‌بندی کن و فقط JSON "
                            'خروجی بده بدون هیچ توضیح اضافه: {"domain": "...", "subdomain": "..."}\n'
                            f"Taxonomy:\n{taxonomy_prompt_text()}"
                        )
                    ),
                    HumanMessage(content=f"منبع: {source}\nمتن:\n{snippet}"),
                ]
            )
            raw = getattr(resp, "content", "") or ""
            m = re.search(r"\{.*\}", raw, re.DOTALL)
            if m:
                parsed = _json.loads(m.group(0))
                tag = {
                    "domain": parsed.get("domain") or "نامشخص",
                    "subdomain": parsed.get("subdomain"),
                    "confidence": 0.75,
                    "method": "llm",
                }
        except Exception:
            pass
    return tag


def _make_document(content: str, metadata: Dict[str, Any]) -> Document:
    return Document(page_content=_e5_passage_text(content), metadata=metadata)


def _load_pdf(path: Path) -> str:
    reader = PdfReader(path.as_posix())
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _load_docx(path: Path) -> str:
    if not _HAS_DOCX:
        raise ImportError(
            "python-docx is not installed. Install with 'pip install python-docx' to load DOCX files."
        )
    doc = DocxDocument(path.as_posix())  # type: ignore
    return "\n".join(p.text for p in doc.paragraphs)


def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _load_docx(path)
    if suffix in {".txt", ".md"}:
        return _load_txt(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _find_legal_units(text: str) -> List[Tuple[int, int, str, str]]:
    """Return list of (start, end, kind, title) ranges for Persian legal units.

    Kinds: ماده, اصل, تبصره, بند. Supports Persian/Latin digits.
    """
    # Heading pattern with optional number/title
    digit = r"[0-9\u06F0-\u06F9]+"
    kinds = ["ماده", "اصل", "تبصره", "بند"]
    # Lookahead split points for headings at line starts or after double newline
    pattern = re.compile(rf"(?m)(?=^(?:{'|'.join(kinds)})\s+{digit}[^\n]*$)")

    # Collect indices of headings
    indices = [m.start() for m in pattern.finditer(text)]
    if not indices:
        return []
    ranges: List[Tuple[int, int, str, str]] = []
    indices.append(len(text))
    for i in range(len(indices) - 1):
        start = indices[i]
        end = indices[i + 1]
        header_line = text[
            start : text.find("\n", start, end) if "\n" in text[start:end] else end
        ]
        kind_match = re.match(
            rf"^(?P<kind>{'|'.join(kinds)})\s+(?P<num>{digit})(?P<rest>[^\n]*)",
            header_line,
        )
        if kind_match:
            kind = kind_match.group("kind")
            title = (kind_match.group("num") + (kind_match.group("rest") or "")).strip()
        else:
            kind = "بخش"
            title = header_line.strip()
        ranges.append((start, end, kind, title))
    return ranges


def _detect_document_type(source: str, text: str) -> str:
    """Detect document type from filename and content."""
    source_lower = source.lower()
    text_lower = text[:500].lower()  # Check first 500 chars

    if "قانون" in source_lower or "law" in source_lower:
        return "law"
    if (
        "آیین‌نامه" in source_lower
        or "regulation" in source_lower
        or "آیین نامه" in source_lower
    ):
        return "regulation"
    if "رای" in source_lower or "ruling" in source_lower or "حکم" in source_lower:
        return "ruling"
    if "قانون" in text_lower:
        return "law"
    if "آیین‌نامه" in text_lower or "آیین نامه" in text_lower:
        return "regulation"
    return "document"


def _detect_legal_domain(text: str) -> str:
    """Detect legal domain from content."""
    text_lower = text[:1000].lower()  # Check first 1000 chars

    domain_keywords = {
        "criminal": ["جرم", "مجازات", "کیفری", "زندان", "حبس"],
        "civil": ["حقوق مدنی", "عقد", "قرارداد", "ارث", "وصیت"],
        "family": ["خانواده", "ازدواج", "طلاق", "نفقه", "حضانت"],
        "commercial": ["تجاری", "شرکت", "سهامی", "چک", "برات"],
    }

    scores = {}
    for domain, keywords in domain_keywords.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > 0:
            scores[domain] = score

    if scores:
        return max(scores.items(), key=lambda x: x[1])[0]
    return "unknown"


def _legal_chunk_documents(text: str, source: str) -> List[Document]:
    """Chunk legal documents with enhanced metadata."""
    document_type = _detect_document_type(source, text)
    legal_domain = _detect_legal_domain(text)
    tax = _tag_taxonomy(source, text)
    domain = tax.get("domain") or "نامشخص"
    subdomain = tax.get("subdomain")

    units = _find_legal_units(text)
    documents: List[Document] = []
    if units:
        for idx, (s, e, kind, title) in enumerate(units):
            content = text[s:e].strip()
            if not content:
                continue
            documents.append(
                _make_document(
                    content,
                    _build_chunk_metadata(
                        source,
                        document_type,
                        legal_domain,
                        content,
                        unit_kind=kind,
                        unit_title=title,
                        unit_index=idx,
                        domain=domain,
                        subdomain=subdomain,
                    ),
                )
            )
        return documents

    # Fallback to general splitter if no legal units detected
    separators = ["\n\n", "\n", "۔", ".", "!", "؟", "?", ";", "،", ",", " "]
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=separators,
        add_start_index=True,
    )
    docs: List[Document] = []
    for doc in splitter.create_documents([text], metadatas=[{}]):
        content = doc.page_content.strip()
        if not content:
            continue
        docs.append(
            _make_document(
                content,
                _build_chunk_metadata(
                    source,
                    document_type,
                    legal_domain,
                    content,
                    unit_index=doc.metadata.get("start_index"),
                    domain=domain,
                    subdomain=subdomain,
                ),
            )
        )
    return docs


def chunk_text(text: str, source: str) -> List[Document]:
    from app.services.content_validation import (
        append_rejection,
        validate_document_content,
    )

    ok, reason = validate_document_content(source, text)
    if not ok:
        append_rejection(
            {
                "source": source,
                "reason": reason,
                "content_chars": len(text or ""),
                "content_head": (text or "")[:240].replace("\n", " "),
            }
        )
        return []
    return _legal_chunk_documents(text, source)


def ingest_files(paths: List[Path]) -> List[Document]:
    documents: List[Document] = []
    for p in paths:
        content = load_text_from_file(p)
        documents.extend(chunk_text(content, source=p.name))
    return documents
