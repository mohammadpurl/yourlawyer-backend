"""Template fill/render services — no LLM, no quota."""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from fastapi import HTTPException
from jinja2 import StrictUndefined
from jinja2.sandbox import SandboxedEnvironment
from sqlalchemy.orm import Session

from app.models.template import DocumentTemplate, TemplateCategory

_env = SandboxedEnvironment(
    undefined=StrictUndefined,
    autoescape=False,
)


def _sanitize_value(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, (int, float, bool)):
        return value
    text = str(value)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = text.replace("{{", "").replace("}}", "").replace("{%", "").replace("%}", "")
    return text


def validate_field_values(field_schema: dict, field_values: dict) -> dict[str, Any]:
    fields = (field_schema or {}).get("fields") or []
    cleaned: dict[str, Any] = {}
    for field in fields:
        key = field.get("key")
        if not key:
            continue
        required = bool(field.get("required"))
        ftype = (field.get("type") or "string").lower()
        raw = field_values.get(key) if field_values else None
        if raw is None or raw == "":
            if required:
                raise HTTPException(
                    status_code=422,
                    detail=f"فیلد الزامی پر نشده است: {field.get('label') or key}",
                )
            continue
        if ftype == "number":
            try:
                cleaned[key] = float(raw) if "." in str(raw) else int(str(raw))
            except (TypeError, ValueError):
                raise HTTPException(
                    status_code=422,
                    detail=f"نوع فیلد نامعتبر است (عدد): {field.get('label') or key}",
                ) from None
        else:
            cleaned[key] = _sanitize_value(raw)
    return cleaned


def render_template_body(body_template: str, field_values: dict[str, Any]) -> str:
    try:
        tmpl = _env.from_string(body_template)
        return tmpl.render(**field_values)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"خطا در جایگذاری قالب: {exc}",
        ) from exc


def build_docx_bytes(title: str, body: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in body.split("\n"):
        doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def list_templates(
    db: Session,
    *,
    category: str | None = None,
    subcategory: str | None = None,
    active_only: bool = True,
) -> list[DocumentTemplate]:
    q = db.query(DocumentTemplate)
    if active_only:
        q = q.filter(DocumentTemplate.is_active.is_(True))
    if category:
        q = q.filter(DocumentTemplate.category == category)
    if subcategory:
        q = q.filter(DocumentTemplate.subcategory == subcategory)
    return q.order_by(
        DocumentTemplate.category,
        DocumentTemplate.subcategory,
        DocumentTemplate.id,
    ).all()


def get_template(
    db: Session, template_id: int, *, active_only: bool = True
) -> DocumentTemplate:
    row = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not row or (active_only and not row.is_active):
        raise HTTPException(status_code=404, detail="قالب یافت نشد")
    return row


def list_categories(db: Session) -> list[dict]:
    rows = (
        db.query(DocumentTemplate.category, DocumentTemplate.subcategory)
        .filter(DocumentTemplate.is_active.is_(True))
        .distinct()
        .all()
    )
    grouped: dict[str, set[str]] = {}
    for cat, sub in rows:
        grouped.setdefault(cat, set()).add(sub)
    return [
        {
            "category": cat,
            "category_label": _category_label(cat),
            "subcategories": sorted(subs),
        }
        for cat, subs in sorted(grouped.items())
    ]


def _category_label(cat: str) -> str:
    return {
        TemplateCategory.CONTRACT.value: "قراردادها",
        TemplateCategory.PRISONER_REQUEST.value: "درخواست‌های استاندارد",
    }.get(cat, cat)


def create_template(db: Session, data: dict) -> DocumentTemplate:
    row = DocumentTemplate(**data)
    db.add(row)
    db.commit()
    db.refresh(row)
    return row


def update_template(db: Session, template_id: int, data: dict) -> DocumentTemplate:
    row = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="قالب یافت نشد")
    for key, value in data.items():
        if value is not None and hasattr(row, key):
            setattr(row, key, value)
    db.commit()
    db.refresh(row)
    return row


def fill_template_docx(
    db: Session, template_id: int, field_values: dict
) -> tuple[bytes, str]:
    tmpl = get_template(db, template_id, active_only=True)
    cleaned = validate_field_values(tmpl.field_schema or {}, field_values or {})
    body = render_template_body(tmpl.body_template, cleaned)
    return build_docx_bytes(tmpl.title, body), f"{tmpl.title}.docx"


def raw_template_docx(db: Session, template_id: int) -> tuple[bytes, str]:
    tmpl = get_template(db, template_id, active_only=True)
    return build_docx_bytes(tmpl.title, tmpl.body_template), f"{tmpl.title}-raw.docx"
